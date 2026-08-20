"""
Skills 自动化测试脚本
无需手动准备文件，程序自动生成测试文档并验证三个Skill
"""

import os
import sys

sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from skills.doc_parser import DocParserSkill
from skills.question_generator import QuestionGeneratorSkill
from skills.solution_optimizer import SolutionOptimizerSkill

PASS = "✓ 通过"
FAIL = "✗ 失败"
results = []


def test(name, condition, detail=""):
    status = PASS if condition else FAIL
    results.append((name, status, detail))
    print(f"  {status} {name}" + (f" — {detail}" if detail else ""))


def create_test_docx(path):
    doc = Document()
    doc.add_heading("大创项目申报书", level=1)
    doc.add_paragraph("项目背景：本项目基于AI知识图谱与多模态视觉技术，构建智能问答系统。")
    doc.add_paragraph("研究目标：实现知识图谱驱动的多模态理解与推理。")
    doc.add_paragraph("创新点：提出跨模态知识对齐方法，融合视觉与语义信息。")
    doc.add_paragraph("技术路线：采用CLIP模型进行图文对齐，结合GNN进行图谱推理。")
    doc.add_paragraph("算法设计：基于Transformer架构，引入Cross-Attention实现多模态融合。")
    doc.add_paragraph("实验评估：在自建数据集上准确率达到92.3%，优于基线方法。")
    doc.add_paragraph("系统架构：前端Vue3 + 后端FastAPI + 知识图谱Neo4j + 向量数据库Milvus。")
    doc.add_paragraph("局限性：当前仅支持图文模态，视频理解有待扩展。")
    doc.add_paragraph("未来规划：扩展视频模态支持，优化实时推理性能。")
    table = doc.add_table(rows=3, cols=3)
    for i in range(3):
        for j in range(3):
            table.rows[i].cells[j].text = f"行{i+1}列{j+1}"
    doc.save(path)


def create_test_pdf(path):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(path, pagesize=letter)
    c.setFont("Helvetica", 12)
    c.drawString(100, 700, "Test PDF for DocParserSkill")
    c.drawString(100, 680, "This is a sample project report for testing.")
    c.save()


print("=" * 60)
print("Skills 自动化测试")
print("=" * 60)

# ========== Skill1: 文档解析 ==========
print("\n【Skill1: 文档解析 DocParserSkill】")

docx_path = os.path.join(os.path.dirname(__file__), "_test_sample.docx")
create_test_docx(docx_path)
test("创建测试docx文件", os.path.exists(docx_path))

parser = DocParserSkill()
text = parser.parse(docx_path)
test("解析docx提取文本", len(text) > 0, f"提取{len(text)}字")
test("文本包含关键内容", "知识图谱" in text and "多模态" in text)

meta = parser.get_metadata()
test("获取元信息", "extension" in meta and meta["extension"] == ".docx")

sections = parser.get_sections(keywords=["创新点", "技术路线"])
test("按章节切分", "创新点" in sections and "技术路线" in sections)

batch = DocParserSkill.batch_parse([docx_path])
test("批量解析", docx_path in batch and len(batch[docx_path]) > 0)

dummy_path = os.path.join(os.path.dirname(__file__), "_test_dummy.xyz")
with open(dummy_path, "w", encoding="utf-8") as f:
    f.write("dummy")
try:
    parser.parse(dummy_path)
    test("不支持的格式报错", False)
except ValueError:
    test("不支持的格式报错", True, "正确抛出ValueError")
os.remove(dummy_path)

try:
    parser.parse("not_exist.docx")
    test("不存在文件报错", False)
except FileNotFoundError:
    test("不存在文件报错", True, "正确抛出FileNotFoundError")

os.remove(docx_path)

# ========== Skill2: 答辩问题生成 ==========
print("\n【Skill2: 答辩问题生成 QuestionGeneratorSkill】")

sample = """
项目背景：基于AI知识图谱与多模态视觉技术构建智能问答系统。
研究目标：实现知识图谱驱动的多模态理解与推理。
创新点：提出跨模态知识对齐方法。
技术路线：采用CLIP模型进行图文对齐，结合GNN进行图谱推理。
算法设计：基于Transformer架构，引入Cross-Attention实现多模态融合。
实验评估：准确率达到92.3%，优于基线方法。
系统架构：前端Vue3 + 后端FastAPI + 知识图谱Neo4j。
局限性：当前仅支持图文模态，视频理解有待扩展。
未来规划：扩展视频模态支持，优化实时推理性能。
"""

gen = QuestionGeneratorSkill()
result = gen.generate(sample)
test("生成答辩问题", True)
test("基础问题非空", len(result.basic_questions) > 0, f"{len(result.basic_questions)}个")
test("技术问题非空", len(result.technical_questions) > 0, f"{len(result.technical_questions)}个")
test("改进问题非空", len(result.improvement_questions) > 0, f"{len(result.improvement_questions)}个")

d = result.to_dict()
test("转字典格式", all(k in d for k in ["基础项目问题", "技术细节问题", "未来改进方向问题"]))

s = result.summary()
test("摘要输出", "基础项目问题" in s and "技术细节问题" in s)

export_path = os.path.join(os.path.dirname(__file__), "_test_questions.txt")
gen.export_questions(export_path)
test("导出文件", os.path.exists(export_path))
os.remove(export_path)

test("get_result返回一致", gen.get_result() is result)

# 空文本兜底
empty_result = gen.generate("")
test("空文本兜底问题", len(empty_result.basic_questions) > 0, "使用默认问题")

# ========== Skill3: 方案优化 ==========
print("\n【Skill3: 方案优化 SolutionOptimizerSkill】")

sample2 = """
本项目构建了基于知识图谱的多模态视觉问答系统。
知识图谱采用Neo4j存储，实体关系通过NLP抽取。
多模态部分使用CLIP模型进行图文对齐，图像特征通过ViT提取。
系统部署为Web服务，提供REST API接口。
前端使用Vue3实现交互式问答界面。
实验表明多模态融合显著提升了问答准确率。
"""

opt = SolutionOptimizerSkill()
opt_result = opt.optimize(sample2)
test("生成优化建议", True)
test("系统改进非空", len(opt_result.system_improvements) > 0, f"{len(opt_result.system_improvements)}条")
test("算法优化非空", len(opt_result.algorithm_optimizations) > 0, f"{len(opt_result.algorithm_optimizations)}条")
test("Demo升级非空", len(opt_result.demo_upgrades) > 0, f"{len(opt_result.demo_upgrades)}条")

d2 = opt_result.to_dict()
test("转字典格式", all(k in d2 for k in ["系统改进思路", "算法优化方向", "Demo升级建议"]))

s2 = opt_result.summary()
test("摘要输出", "系统改进思路" in s2 and "算法优化方向" in s2)

export_path2 = os.path.join(os.path.dirname(__file__), "_test_optimizations.txt")
opt.export_suggestions(export_path2)
test("导出文件", os.path.exists(export_path2))
os.remove(export_path2)

test("get_result返回一致", opt.get_result() is opt_result)

# 空文本兜底
empty_opt = opt.optimize("无关键词文本")
test("无匹配时兜底建议", len(empty_opt.system_improvements) > 0, "使用通用建议")

# ========== 组合流水线 ==========
print("\n【组合流水线: 解析→问题生成+方案优化】")

docx_path2 = os.path.join(os.path.dirname(__file__), "_test_pipeline.docx")
create_test_docx(docx_path2)

parser2 = DocParserSkill()
pipeline_text = parser2.parse(docx_path2)
test("流水线: 解析文档", len(pipeline_text) > 0)

q_result = gen.generate(pipeline_text)
test("流水线: 生成问题", len(q_result.basic_questions) > 0)

o_result = opt.optimize(pipeline_text)
test("流水线: 生成优化", len(o_result.system_improvements) > 0)

os.remove(docx_path2)

# ========== 汇总 ==========
print("\n" + "=" * 60)
passed = sum(1 for _, s, _ in results if s == PASS)
failed = sum(1 for _, s, _ in results if s == FAIL)
print(f"测试结果: {passed} 通过 / {failed} 失败 / 共 {len(results)} 项")
if failed > 0:
    print("\n失败项:")
    for name, s, d in results:
        if s == FAIL:
            print(f"  - {name} {d}")
print("=" * 60)