"""
Skill1: 文档解析Skill
支持读取docx、pdf文件，提取大创项目申报书、论文、项目报告全文文本。
"""

import os
from typing import Optional
from docx import Document
from PyPDF2 import PdfReader


class DocParserSkill:
    """文档解析Skill，支持docx和pdf格式"""

    SUPPORTED_EXTENSIONS = {".docx", ".pdf"}

    def __init__(self):
        self._raw_text: str = ""
        self._metadata: dict = {}

    def parse(self, file_path: str) -> str:
        """
        解析文档，自动识别格式并提取全文文本。

        Args:
            file_path: 文件路径（支持.docx / .pdf）

        Returns:
            提取的全文文本

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件格式
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}，仅支持 {self.SUPPORTED_EXTENSIONS}")

        if ext == ".docx":
            self._raw_text = self._parse_docx(file_path)
        elif ext == ".pdf":
            self._raw_text = self._parse_pdf(file_path)

        self._metadata["file_path"] = file_path
        self._metadata["extension"] = ext
        self._metadata["char_count"] = len(self._raw_text)
        return self._raw_text

    def _parse_docx(self, file_path: str) -> str:
        """解析docx文件，提取所有段落文本"""
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    paragraphs.append(row_text)

        return "\n".join(paragraphs)

    def _parse_pdf(self, file_path: str) -> str:
        """解析pdf文件，提取所有页面文本"""
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n".join(pages)

    def get_metadata(self) -> dict:
        """获取文档元信息"""
        return self._metadata.copy()

    def get_sections(self, keywords: Optional[list] = None) -> dict:
        """
        按关键词切分文本为不同章节。

        Args:
            keywords: 章节标题关键词列表，如 ["摘要", "技术路线", "创新点"]

        Returns:
            {关键词: 对应文本} 的字典
        """
        if not keywords:
            return {"full_text": self._raw_text}

        sections = {}
        lines = self._raw_text.split("\n")
        current_key = "其他"
        sections[current_key] = []

        for line in lines:
            for kw in keywords:
                if kw in line:
                    current_key = kw
                    if current_key not in sections:
                        sections[current_key] = []
                    break
            sections.setdefault(current_key, []).append(line)

        return {k: "\n".join(v).strip() for k, v in sections.items() if v}

    @staticmethod
    def batch_parse(file_paths: list) -> dict:
        """
        批量解析多个文档。

        Args:
            file_paths: 文件路径列表

        Returns:
            {文件路径: 提取文本} 的字典
        """
        results = {}
        parser = DocParserSkill()
        for fp in file_paths:
            try:
                results[fp] = parser.parse(fp)
            except Exception as e:
                results[fp] = f"[解析失败] {e}"
        return results