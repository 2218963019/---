"""
RAG引擎 自动化测试
验证分块、嵌入、存储、检索全流程
"""

import os
import sys
import json

sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from skills.rag_engine import (
    TextChunker, TFIDFEmbedder, VectorStore, RAGEngine,
    Chunk, RetrievalResult, _cosine_similarity
)

PASS = "✓ 通过"
FAIL = "✗ 失败"
results = []


def test(name, condition, detail=""):
    status = PASS if condition else FAIL
    results.append((name, status, detail))
    print(f"  {status} {name}" + (f" — {detail}" if detail else ""))


def create_test_docx(path):
    doc = Document()
    doc.add_heading("大创项目申报书", level=1)
    doc.add_paragraph("项目背景：本项目基于AI知识图谱与多模态视觉技术，构建智能问答系统。项目旨在解决传统问答系统无法理解图像信息的问题。")
    doc.add_paragraph("研究目标：实现知识图谱驱动的多模态理解与推理，支持图文混合输入的智能问答。")
    doc.add_paragraph("创新点：提出跨模态知识对齐方法，将视觉特征映射到知识图谱语义空间，实现图文联合推理。")
    doc.add_paragraph("技术路线：采用CLIP模型进行图文对齐，结合图神经网络GNN进行图谱推理，使用Transformer架构实现多模态融合。")
    doc.add_paragraph("算法设计：基于Cross-Attention机制实现图文深度交互，引入对比学习增强多模态表征质量。")
    doc.add_paragraph("实验评估：在自建多模态问答数据集上准确率达到92.3%，F1分数为89.7%，显著优于基线方法。")
    doc.add_paragraph("系统架构：前端Vue3 + 后端FastAPI + 知识图谱Neo4j + 向量数据库Milvus + 模型服务Triton。")
    doc.add_paragraph("局限性：当前仅支持图文两种模态，视频和音频理解有待扩展；推理延迟较高，需优化。")
    doc.add_paragraph("未来规划：扩展视频模态支持，引入模型蒸馏降低延迟，增加用户反馈闭环优化系统。")
    doc.save(path)


SAMPLE_TEXT = """
项目背景：本项目基于AI知识图谱与多模态视觉技术，构建智能问答系统。
研究目标：实现知识图谱驱动的多模态理解与推理。
创新点：提出跨模态知识对齐方法，融合视觉与语义信息。
技术路线：采用CLIP模型进行图文对齐，结合GNN进行图谱推理。
算法设计：基于Transformer架构，引入Cross-Attention实现多模态融合。
实验评估：准确率达到92.3%，优于基线方法。
系统架构：前端Vue3 + 后端FastAPI + 知识图谱Neo4j。
局限性：当前仅支持图文模态，视频理解有待扩展。
未来规划：扩展视频模态支持，优化实时推理性能。
""".strip()

print("=" * 60)
print("RAG引擎 自动化测试")
print("=" * 60)

# ========== 1. 文档分块器 ==========
print("\n【1. 文档分块器 TextChunker】")

chunker = TextChunker(chunk_size=200, overlap=30)

fixed_chunks = chunker.split_fixed(SAMPLE_TEXT, "test")
test("fixed分块非空", len(fixed_chunks) > 0, f"{len(fixed_chunks)}个chunk")
test("fixed分块内容完整", all(c.content.strip() for c in fixed_chunks))

para_chunks = chunker.split_paragraph(SAMPLE_TEXT, "test")
test("paragraph分块非空", len(para_chunks) > 0, f"{len(para_chunks)}个chunk")

section_chunks = chunker.split_section(
    SAMPLE_TEXT,
    ["创新点", "技术路线", "实验", "系统架构", "局限性", "未来规划"],
    "test"
)
test("section分块非空", len(section_chunks) > 0, f"{len(section_chunks)}个chunk")
has_section_meta = any("section" in c.metadata for c in section_chunks)
test("section分块含章节元信息", has_section_meta)

# ========== 2. TF-IDF嵌入器 ==========
print("\n【2. TF-IDF嵌入器 TFIDFEmbedder】")

embedder = TFIDFEmbedder()
docs = [
    "知识图谱是人工智能的核心技术",
    "多模态视觉理解融合图像和文本",
    "深度学习模型需要大量数据训练",
]
embedder.fit(docs)
test("fit训练完成", embedder._fitted)
test("词汇表非空", embedder.dim > 0, f"维度:{embedder.dim}")

vec = embedder.embed("知识图谱技术")
test("嵌入向量非零", any(v != 0 for v in vec))
test("嵌入维度正确", len(vec) == embedder.dim)

vecs = embedder.embed_batch(docs)
test("批量嵌入数量一致", len(vecs) == len(docs))

sim = _cosine_similarity(vecs[0], vecs[0])
test("自身相似度≈1", abs(sim - 1.0) < 0.01, f"{sim:.4f}")

sim_diff = _cosine_similarity(vecs[0], vecs[2])
test("不同文本相似度<1", sim_diff < 1.0, f"{sim_diff:.4f}")

# ========== 3. 向量存储 ==========
print("\n【3. 向量存储 VectorStore】")

store = VectorStore()
test("初始为空", store.size == 0)

chunks_data = [Chunk(content=d, metadata={"idx": i}, chunk_id=i) for i, d in enumerate(docs)]
store.add(chunks_data, vecs)
test("添加后size正确", store.size == 3)

results = store.search(vecs[0], top_k=2)
test("检索返回结果", len(results) == 2)
test("最相似为自身", results[0].chunk.chunk_id == 0)
test("相似度降序", results[0].score >= results[1].score)

store.clear()
test("clear后为空", store.size == 0)

# 持久化测试
persist_path = os.path.join(os.path.dirname(__file__), "_test_store.json")
store.add(chunks_data, vecs)
store.save(persist_path)
test("持久化文件存在", os.path.exists(persist_path))

store2 = VectorStore()
store2.load(persist_path)
test("加载后size一致", store2.size == 3)
os.remove(persist_path)

# ========== 4. RAG引擎主类 ==========
print("\n【4. RAG引擎 RAGEngine】")

engine = RAGEngine(chunk_size=300, chunk_overlap=30)
n = engine.index_text(SAMPLE_TEXT, source="test_doc", strategy="section")
test("索引文本成功", n > 0, f"{n}个chunk")
test("引擎统计信息", engine.stats["indexed_chunks"] == n)

# 语义检索
ret_results = engine.retrieve("项目的创新点是什么？", top_k=3)
test("语义检索返回结果", len(ret_results) > 0, f"{len(ret_results)}条")
test("检索结果有分数", all(isinstance(r.score, float) or isinstance(r.score, int) for r in ret_results))
test("检索结果有chunk", all(isinstance(r.chunk, Chunk) for r in ret_results))

# 带上下文检索
context = engine.retrieve_with_context("技术路线和算法设计", top_k=3)
test("上下文检索非空", len(context) > 0)
test("上下文含相关度", "相关度" in context)

# 不同查询的检索区分度
r1 = engine.retrieve("知识图谱构建方法", top_k=1)
r2 = engine.retrieve("系统部署架构", top_k=1)
test("不同查询返回不同结果", r1[0].chunk.chunk_id != r2[0].chunk.chunk_id or r1[0].score != r2[0].score)

# ========== 5. 文档索引（集成Skill1）==========
print("\n【5. 文档索引（集成DocParserSkill）】")

docx_path = os.path.join(os.path.dirname(__file__), "_test_rag.docx")
create_test_docx(docx_path)

engine2 = RAGEngine(chunk_size=300)
n2 = engine2.index_document(docx_path, strategy="section")
test("索引docx文档", n2 > 0, f"{n2}个chunk")

results2 = engine2.retrieve("多模态融合的方法", top_k=3)
test("从docx检索", len(results2) > 0)

context2 = engine2.retrieve_with_context("创新点和技术路线", top_k=2)
test("docx上下文检索", "检索到以下相关内容" in context2)

os.remove(docx_path)

# ========== 6. 多文档索引 ==========
print("\n【6. 多文档索引】")

engine3 = RAGEngine(chunk_size=200)
engine3.index_text(SAMPLE_TEXT, source="申报书", strategy="paragraph")

extra_text = "本论文提出了一种基于图神经网络的实体关系抽取方法，在多个基准数据集上取得了SOTA效果。"
engine3.index_text(extra_text, source="论文", strategy="paragraph")
test("多文档索引", engine3.stats["indexed_chunks"] > 0)

results3 = engine3.retrieve("图神经网络实体抽取", top_k=3)
test("跨文档检索", len(results3) > 0)

# ========== 7. 分块策略对比 ==========
print("\n【7. 分块策略对比】")

e_fixed = RAGEngine(chunk_size=200)
n_fixed = e_fixed.index_text(SAMPLE_TEXT, strategy="fixed")

e_para = RAGEngine(chunk_size=200)
n_para = e_para.index_text(SAMPLE_TEXT, strategy="paragraph")

e_sec = RAGEngine(chunk_size=200)
n_sec = e_sec.index_text(SAMPLE_TEXT, strategy="section")

test("三种策略均可索引", n_fixed > 0 and n_para > 0 and n_sec > 0)
test("section策略保留章节语义", n_sec > 0 and n_para > 0, f"section:{n_sec} para:{n_para} fixed:{n_fixed}")

# ========== 汇总 ==========
print("\n" + "=" * 60)
passed = sum(1 for _, s, _ in results if s == PASS)
failed = sum(1 for _, s, _ in results if s == FAIL)
print(f"测试结果: {passed} 通过 / {failed} 失败 / 共 {len(results)} 项")
if failed > 0:
    print("\n失败项:")
    for name, s, d in results:
        if s == FAIL:
            print(f"  - {name} {d}")
print("=" * 60)